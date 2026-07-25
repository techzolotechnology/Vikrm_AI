import io

import pytest

from app.services.rag.parsers import (
    DocumentParseError,
    UnsupportedFileTypeError,
    parse_document,
)


def test_parse_txt() -> None:
    text = parse_document(filename="notes.txt", content=b"Hello, this is a plain text file.")
    assert text == "Hello, this is a plain text file."


def test_parse_markdown() -> None:
    text = parse_document(filename="readme.md", content=b"# Title\n\nSome **markdown** content.")
    assert "Title" in text
    assert "markdown" in text


def test_parse_csv() -> None:
    csv_bytes = b"name,age\nAlice,30\nBob,25\n"
    text = parse_document(filename="people.csv", content=csv_bytes)
    assert "Alice, 30" in text
    assert "Bob, 25" in text


def test_parse_docx() -> None:
    from docx import Document as DocxDocument

    buffer = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("This is a real Word document paragraph.")
    doc.add_paragraph("A second paragraph with more content.")
    doc.save(buffer)

    text = parse_document(filename="report.docx", content=buffer.getvalue())
    assert "real Word document paragraph" in text
    assert "second paragraph" in text


def test_parse_pdf() -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "This text lives inside a real generated PDF.")
    pdf_bytes = bytes(pdf.output())

    text = parse_document(filename="doc.pdf", content=pdf_bytes)
    assert "real generated PDF" in text


def test_unsupported_extension_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        parse_document(filename="archive.zip", content=b"PK\x03\x04")


def test_invalid_utf8_text_raises() -> None:
    with pytest.raises(DocumentParseError):
        parse_document(filename="broken.txt", content=b"\xff\xfe\x00invalid")


def test_no_extension_raises_unsupported() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        parse_document(filename="README", content=b"some content")
